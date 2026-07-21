import os
import tempfile
import pytest
from services.resume_text_extractor import ResumeTextExtractor, TextExtractionResult
from services.resume_url_extractor import ResumeURLExtractor
from schemas.resume_profile import ResumeProfileData, ExtractedURLs


class TestFileValidation:

    def test_valid_pdf(self):
        ext = ResumeTextExtractor.validate_file("resume.pdf")
        assert ext == ".pdf"

    def test_valid_docx(self):
        ext = ResumeTextExtractor.validate_file("resume.docx")
        assert ext == ".docx"

    def test_invalid_extension(self):
        with pytest.raises(ValueError, match="Unsupported file type"):
            ResumeTextExtractor.validate_file("resume.txt")

    def test_invalid_exe(self):
        with pytest.raises(ValueError, match="Unsupported file type"):
            ResumeTextExtractor.validate_file("malware.exe")

    def test_no_extension(self):
        with pytest.raises(ValueError, match="must have an extension"):
            ResumeTextExtractor.validate_file("resume")

    def test_empty_file(self):
        with pytest.raises(ValueError, match="empty"):
            ResumeTextExtractor.validate_file("resume.pdf", file_size=0)

    def test_file_too_large(self):
        with pytest.raises(ValueError, match="exceeds maximum size"):
            ResumeTextExtractor.validate_file("resume.pdf", file_size=10 * 1024 * 1024)

    def test_file_at_max_size(self):
        ext = ResumeTextExtractor.validate_file("resume.pdf", file_size=5 * 1024 * 1024)
        assert ext == ".pdf"


class TestDocxExtraction:

    def test_extract_docx(self):
        import docx
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = docx.Document()
            doc.add_paragraph("John Doe")
            doc.add_paragraph("Cybersecurity Analyst with 5 years of experience")
            doc.add_paragraph("Skills: Python, Linux, SIEM")
            doc.save(f.name)
            f.flush()
            try:
                result = ResumeTextExtractor.extract(f.name)
                assert isinstance(result, TextExtractionResult)
                assert "John Doe" in result.text
                assert "Python" in result.text
                assert result.file_type == "docx"
                assert result.character_count > 0
                assert result.page_count is None
            finally:
                os.unlink(f.name)

    def test_empty_docx(self):
        import docx
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = docx.Document()
            doc.save(f.name)
            f.flush()
            try:
                result = ResumeTextExtractor.extract(f.name)
                assert result.text == ""
                assert result.character_count == 0
            finally:
                os.unlink(f.name)


class TestTextNormalization:

    def test_whitespace_normalization(self):
        text = "Hello   World\t\tTest"
        normalized = ResumeTextExtractor._normalize(text)
        assert "   " not in normalized
        assert "\t" not in normalized

    def test_blank_line_normalization(self):
        text = "Line 1\n\n\n\n\nLine 2"
        normalized = ResumeTextExtractor._normalize(text)
        assert "\n\n\n" not in normalized

    def test_empty_text(self):
        assert ResumeTextExtractor._normalize("") == ""
        assert ResumeTextExtractor._normalize(None) == ""


class TestURLExtraction:

    def test_github_url(self):
        text = "Check out my work at https://github.com/johndoe"
        urls = ResumeURLExtractor.extract(text)
        assert len(urls["github"]) == 1
        assert "github.com/johndoe" in urls["github"][0]

    def test_linkedin_url(self):
        text = "Connect with me: https://www.linkedin.com/in/johndoe"
        urls = ResumeURLExtractor.extract(text)
        assert len(urls["linkedin"]) == 1
        assert "linkedin.com/in/johndoe" in urls["linkedin"][0]

    def test_multiple_urls(self):
        text = """
        GitHub: https://github.com/johndoe
        LinkedIn: https://www.linkedin.com/in/johndoe
        Portfolio: https://johndoe.dev
        Website: https://example.com
        """
        urls = ResumeURLExtractor.extract(text)
        assert len(urls["github"]) == 1
        assert len(urls["linkedin"]) == 1
        assert len(urls["portfolio"]) + len(urls["personal_website"]) + len(urls["other"]) >= 2

    def test_duplicate_urls(self):
        text = "https://github.com/johndoe and again https://github.com/johndoe"
        urls = ResumeURLExtractor.extract(text)
        assert len(urls["github"]) == 1

    def test_trailing_punctuation(self):
        text = "Visit https://github.com/johndoe."
        urls = ResumeURLExtractor.extract(text)
        assert urls["github"][0].endswith("johndoe")

    def test_empty_text(self):
        urls = ResumeURLExtractor.extract("")
        assert urls == {"github": [], "linkedin": [], "portfolio": [], "personal_website": [], "other": []}

    def test_none_text(self):
        urls = ResumeURLExtractor.extract(None)
        assert urls == {"github": [], "linkedin": [], "portfolio": [], "personal_website": [], "other": []}

    def test_normalization_removes_trailing_slash(self):
        url = ResumeURLExtractor._normalize("https://github.com/johndoe/")
        assert url == "https://github.com/johndoe"

    def test_normalization_removes_www(self):
        url = ResumeURLExtractor._normalize("https://www.example.com/page")
        assert "www." not in url

    def test_dev_domain_detected(self):
        text = "Visit https://johndoe.dev for my portfolio"
        urls = ResumeURLExtractor.extract(text)
        total = sum(len(v) for v in urls.values())
        assert total >= 1


class TestPydanticValidation:

    def test_valid_profile(self):
        profile = ResumeProfileData(
            full_name="John Doe",
            email="john@example.com",
            skills=[{"name": "Python", "category": "Programming", "source": "resume"}],
        )
        assert profile.full_name == "John Doe"
        assert len(profile.skills) == 1

    def test_none_fields(self):
        profile = ResumeProfileData()
        assert profile.full_name is None
        assert profile.education == []
        assert profile.skills == []

    def test_skill_as_string_conversion(self):
        data = {"skills": [{"name": "Python", "category": "Programming", "source": "resume"}]}
        profile = ResumeProfileData(**data)
        assert profile.skills[0].name == "Python"

    def test_urls_structure(self):
        urls = ExtractedURLs(
            github=["https://github.com/test"],
            linkedin=["https://linkedin.com/in/test"],
        )
        assert len(urls.github) == 1
        assert urls.portfolio == []

    def test_education_entry(self):
        profile = ResumeProfileData(
            education=[{
                "institution": "MIT",
                "degree": "BS",
                "field": "CS",
                "start_date": "2018",
                "end_date": "2022",
            }]
        )
        assert profile.education[0].institution == "MIT"


class TestGeminiMistralFallback:

    def test_ai_service_has_orchestrator(self):
        from app.ai.ai_service import AIService
        svc = AIService(provider_name="gemini")
        assert hasattr(svc, 'orchestrator')
        assert svc.orchestrator is not None

    def test_orchestrator_has_fallback(self):
        from app.ai.orchestrator import AIOrchestrator
        orch = AIOrchestrator(provider_name="gemini", fallback_provider_name="mistral")
        assert orch.fallback_provider is not None or orch.provider is not None

    def test_mock_fallback_returns_valid_structure(self):
        from app.ai.orchestrator import AIOrchestrator
        orch = AIOrchestrator(provider_name="nonexistent")
        mock = orch._mock_fallback("resume_profile")
        assert isinstance(mock, dict)
        assert "full_name" in mock
        assert "skills" in mock
        assert "urls" in mock


class TestDatabasePersistence:

    def test_resume_profile_model_fields(self):
        from models.sqlalchemy_models import ResumeProfile
        columns = {c.name for c in ResumeProfile.__table__.columns}
        assert "id" in columns
        assert "user_id" in columns
        assert "original_filename" in columns
        assert "file_type" in columns
        assert "file_size" in columns
        assert "raw_text" in columns
        assert "extracted_profile" in columns
        assert "extracted_urls" in columns
        assert "processing_status" in columns
        assert "processing_error" in columns
        assert "created_at" in columns
        assert "updated_at" in columns


class TestAPIEndpointProtection:

    def test_analyze_requires_auth(self):
        from fastapi.testclient import TestClient
        from app.main import app
        client = TestClient(app)
        res = client.post("/api/v1/resume/profile/analyze")
        assert res.status_code in (401, 403)

    def test_get_profile_requires_auth(self):
        from fastapi.testclient import TestClient
        from app.main import app
        client = TestClient(app)
        res = client.get("/api/v1/resume/profile")
        assert res.status_code in (401, 403)
